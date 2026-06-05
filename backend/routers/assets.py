from __future__ import annotations

import asyncio
import io
import json
import logging
import re
import urllib.error
import urllib.request
from pathlib import Path
from typing import Any

from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from pydantic import BaseModel, Field

from ..config import DEEPSEEK_API_KEY, DEEPSEEK_BASE_URL, DEEPSEEK_MODEL

router = APIRouter(prefix="/api/assets", tags=["assets"])
log = logging.getLogger("assets")

MAX_TEXT_CHARS = 60000
MAX_FILE_BYTES = 12 * 1024 * 1024


class AnalyzeTextRequest(BaseModel):
    text: str = Field(..., min_length=20, max_length=MAX_TEXT_CHARS)
    template_id: str | None = None
    product_hint: str | None = None
    origin_hint: str | None = None
    file_names: list[str] = Field(default_factory=list)


class PosterChatMessage(BaseModel):
    role: str
    content: str = Field(..., min_length=1, max_length=4000)


class PosterChatRequest(BaseModel):
    messages: list[PosterChatMessage] = Field(default_factory=list, max_length=20)
    poster: dict[str, Any] = Field(default_factory=dict)
    product: dict[str, Any] = Field(default_factory=dict)


class StudioChatRequest(BaseModel):
    messages: list[PosterChatMessage] = Field(default_factory=list, max_length=20)
    output_type: str = Field(..., pattern="^(poster|archive|display)$")
    output: dict[str, Any] = Field(default_factory=dict)
    product: dict[str, Any] = Field(default_factory=dict)
    asset_package: dict[str, Any] = Field(default_factory=dict)


class AssetChatRequest(BaseModel):
    messages: list[PosterChatMessage] = Field(default_factory=list, max_length=20)
    asset_package: dict[str, Any] = Field(default_factory=dict)
    extraction: dict[str, Any] = Field(default_factory=dict)


class AssetAnalysisResponse(BaseModel):
    model: str
    extraction: dict[str, Any]
    thinking_trace: list[str]
    asset_package: dict[str, Any]


class PosterChatResponse(BaseModel):
    model: str
    reply: str
    suggestions: list[dict[str, Any]]
    style_options: list[dict[str, Any]]


class StudioChatResponse(BaseModel):
    model: str
    reply: str
    suggestions: list[dict[str, Any]]
    style_options: list[dict[str, Any]]
    next_actions: list[str]


class AssetChatResponse(BaseModel):
    model: str
    reply: str
    recommendations: list[dict[str, Any]]
    next_actions: list[str]


@router.post("/analyze-text", response_model=AssetAnalysisResponse)
async def analyze_text(payload: AnalyzeTextRequest):
    extracted = {
        "source": "manual_text",
        "files": [
            {
                "name": name,
                "type": "context",
                "status": "referenced",
                "chars": 0,
                "note": "User supplied this file name as context.",
            }
            for name in payload.file_names
        ],
        "combined_chars": len(payload.text),
    }
    result = await _run_asset_analysis(
        text=payload.text,
        extraction=extracted,
        template_id=payload.template_id,
        product_hint=payload.product_hint,
        origin_hint=payload.origin_hint,
    )
    return AssetAnalysisResponse(**result)


@router.post("/analyze-files", response_model=AssetAnalysisResponse)
async def analyze_files(
    files: list[UploadFile] = File(default=[]),
    manual_text: str = Form(default=""),
    template_id: str | None = Form(default=None),
    product_hint: str | None = Form(default=None),
    origin_hint: str | None = Form(default=None),
):
    extraction = await _extract_uploads(files)
    chunks = [
        f"[[source:{item.get('name', 'unknown')}]]\n{item['text']}"
        for item in extraction["files"]
        if item.get("text")
    ]
    if manual_text.strip():
        chunks.append(f"[[source:manual-note.txt]]\n{manual_text.strip()}")
        extraction["files"].append({
            "name": "manual-note.txt",
            "type": "manual_text",
            "status": "extracted",
            "chars": len(manual_text.strip()),
            "note": "User supplied manual text.",
        })

    combined_text = "\n\n---\n\n".join(chunks).strip()
    if len(combined_text) < 20:
        raise HTTPException(
            status_code=422,
            detail="没有提取到足够文本。请上传 PDF/DOCX/XLSX/CSV/TXT/JSON/Markdown，或补充手动说明。",
        )

    combined_text = combined_text[:MAX_TEXT_CHARS]
    extraction["combined_chars"] = len(combined_text)
    result = await _run_asset_analysis(
        text=combined_text,
        extraction=extraction,
        template_id=template_id,
        product_hint=product_hint,
        origin_hint=origin_hint,
    )
    return AssetAnalysisResponse(**result)


@router.post("/poster-chat", response_model=PosterChatResponse)
async def poster_chat(payload: PosterChatRequest):
    _require_deepseek()
    if not payload.messages:
        raise HTTPException(status_code=422, detail="消息不能为空")

    try:
        raw = await asyncio.to_thread(_call_deepseek_json, _poster_chat_messages(payload), 2200)
    except urllib.error.HTTPError as exc:
        body = exc.read().decode("utf-8", errors="replace")
        log.warning("DeepSeek HTTP error %s: %s", exc.code, body[:400])
        raise HTTPException(status_code=502, detail=f"DeepSeek request failed: HTTP {exc.code}") from exc
    except urllib.error.URLError as exc:
        log.warning("DeepSeek network error: %s", exc.reason)
        raise HTTPException(status_code=502, detail="DeepSeek request failed: network error") from exc
    except ValueError as exc:
        raise HTTPException(status_code=502, detail=str(exc)) from exc

    return PosterChatResponse(
        model=DEEPSEEK_MODEL,
        reply=str(raw.get("reply") or "我已经根据当前海报内容给出建议。"),
        suggestions=_normalize_suggestions(raw.get("suggestions")),
        style_options=_as_list_of_dict(raw.get("style_options")),
    )


@router.post("/studio-chat", response_model=StudioChatResponse)
async def studio_chat(payload: StudioChatRequest):
    _require_deepseek()
    if not payload.messages:
        raise HTTPException(status_code=422, detail="消息不能为空")

    try:
        raw = await asyncio.to_thread(_call_deepseek_json, _studio_chat_messages(payload), 2600)
    except urllib.error.HTTPError as exc:
        body = exc.read().decode("utf-8", errors="replace")
        log.warning("DeepSeek HTTP error %s: %s", exc.code, body[:400])
        raise HTTPException(status_code=502, detail=f"DeepSeek request failed: HTTP {exc.code}") from exc
    except urllib.error.URLError as exc:
        log.warning("DeepSeek network error: %s", exc.reason)
        raise HTTPException(status_code=502, detail="DeepSeek request failed: network error") from exc
    except ValueError as exc:
        raise HTTPException(status_code=502, detail=str(exc)) from exc

    return StudioChatResponse(
        model=DEEPSEEK_MODEL,
        reply=str(raw.get("reply") or "我已经根据当前项目内容给出建议。"),
        suggestions=_normalize_studio_suggestions(payload.output_type, raw.get("suggestions")),
        style_options=_as_list_of_dict(raw.get("style_options"))[:8],
        next_actions=_as_text_list(raw.get("next_actions"))[:8],
    )


@router.post("/asset-chat", response_model=AssetChatResponse)
async def asset_chat(payload: AssetChatRequest):
    _require_deepseek()
    if not payload.messages:
        raise HTTPException(status_code=422, detail="消息不能为空")

    try:
        raw = await asyncio.to_thread(_call_deepseek_json, _asset_chat_messages(payload), 2600)
    except urllib.error.HTTPError as exc:
        body = exc.read().decode("utf-8", errors="replace")
        log.warning("DeepSeek HTTP error %s: %s", exc.code, body[:400])
        raise HTTPException(status_code=502, detail=f"DeepSeek request failed: HTTP {exc.code}") from exc
    except urllib.error.URLError as exc:
        log.warning("DeepSeek network error: %s", exc.reason)
        raise HTTPException(status_code=502, detail="DeepSeek request failed: network error") from exc
    except ValueError as exc:
        raise HTTPException(status_code=502, detail=str(exc)) from exc

    return AssetChatResponse(
        model=DEEPSEEK_MODEL,
        reply=str(raw.get("reply") or "我已经根据当前资产包给出下一步建议。"),
        recommendations=_as_list_of_dict(raw.get("recommendations"))[:8],
        next_actions=_as_text_list(raw.get("next_actions"))[:8],
    )


async def _extract_uploads(files: list[UploadFile]) -> dict[str, Any]:
    result: dict[str, Any] = {"source": "uploaded_files", "files": []}
    for upload in files:
        name = Path(upload.filename or "untitled").name
        suffix = Path(name).suffix.lower()
        data = await upload.read()
        if len(data) > MAX_FILE_BYTES:
            result["files"].append(_file_result(name, suffix, "skipped", "", "文件超过 12MB，已跳过。"))
            continue
        try:
            text, note = _extract_text(name, suffix, data)
            status = "extracted" if text else "metadata_only"
            result["files"].append(_file_result(name, suffix, status, text, note))
        except Exception as exc:
            log.warning("File extraction failed for %s: %s", name, exc)
            result["files"].append(_file_result(name, suffix, "failed", "", str(exc)))
    return result


def _extract_text(name: str, suffix: str, data: bytes) -> tuple[str, str]:
    if suffix in {".txt", ".md", ".markdown", ".csv", ".json"}:
        return _decode_text(data), "文本文件已直接读取。"
    if suffix == ".pdf":
        return _extract_pdf(data), "PDF 页面文本已提取。"
    if suffix == ".docx":
        return _extract_docx(data), "DOCX 段落和表格文本已提取。"
    if suffix == ".xlsx":
        return _extract_xlsx(data), "XLSX 前若干工作表单元格已提取。"
    if suffix == ".xls":
        return "", "暂不支持旧版 .xls，请另存为 .xlsx 或 CSV。"
    if suffix in {".png", ".jpg", ".jpeg", ".webp"}:
        return "", "图片已记录为视觉素材；当前未接入 OCR。"
    return _decode_text(data), f"{name} 按文本尝试读取。"


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return data.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace").strip()


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for index, page in enumerate(reader.pages[:30], start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[PDF page {index}]\n{text.strip()}")
    return "\n\n".join(pages).strip()


def _extract_docx(data: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(data))
    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables[:10]:
        for row in table.rows[:80]:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                parts.append(" | ".join(values))
    return "\n".join(parts).strip()


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for sheet in workbook.worksheets[:6]:
        parts.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(max_row=120, values_only=True):
            values = [str(value).strip() for value in row if value not in (None, "")]
            if values:
                parts.append(" | ".join(values))
    return "\n".join(parts).strip()


def _file_result(name: str, suffix: str, status: str, text: str, note: str) -> dict[str, Any]:
    item = {
        "name": name,
        "type": suffix.lstrip(".") or "unknown",
        "status": status,
        "chars": len(text),
        "note": note,
    }
    if text:
        item["text"] = text[:20000]
        item["preview"] = re.sub(r"\s+", " ", text[:240]).strip()
    return item


async def _run_asset_analysis(
    text: str,
    extraction: dict[str, Any],
    template_id: str | None,
    product_hint: str | None,
    origin_hint: str | None,
) -> dict[str, Any]:
    _require_deepseek()
    try:
        raw = await asyncio.to_thread(
            _call_deepseek_json,
            _asset_analysis_messages(text, extraction, template_id, product_hint, origin_hint),
            3600,
        )
    except urllib.error.HTTPError as exc:
        body = exc.read().decode("utf-8", errors="replace")
        log.warning("DeepSeek HTTP error %s: %s", exc.code, body[:400])
        raise HTTPException(status_code=502, detail=f"DeepSeek request failed: HTTP {exc.code}") from exc
    except urllib.error.URLError as exc:
        log.warning("DeepSeek network error: %s", exc.reason)
        raise HTTPException(status_code=502, detail="DeepSeek request failed: network error") from exc
    except ValueError as exc:
        raise HTTPException(status_code=502, detail=str(exc)) from exc

    return {
        "model": DEEPSEEK_MODEL,
        "extraction": _public_extraction(extraction),
        "thinking_trace": _as_text_list(raw.get("thinking_trace"))[:8],
        "asset_package": _normalize_package(raw.get("asset_package") or raw),
    }


def _asset_analysis_messages(
    text: str,
    extraction: dict[str, Any],
    template_id: str | None,
    product_hint: str | None,
    origin_hint: str | None,
) -> list[dict[str, str]]:
    file_summary = [
        {
            "name": item.get("name"),
            "type": item.get("type"),
            "status": item.get("status"),
            "chars": item.get("chars"),
            "note": item.get("note"),
        }
        for item in extraction.get("files", [])
    ]
    schema = {
        "thinking_trace": ["外显分析步骤，简洁说明你如何处理资料，不要输出隐藏推理。"],
        "asset_package": {
            "product": {"name": "", "category": "", "origin": "", "ingredients": []},
            "evidence": {"lab_indicators": [], "certifications": [], "origin_claims": [], "process_steps": []},
            "visualization": {
                "map_nodes": [
                    {
                        "name": "空间节点名称",
                        "short": "短名称",
                        "coord": [0, 0],
                        "desc": "节点在产地、加工、仓储、销售或传播中的作用；coord 必须是 [经度, 纬度]。如果无法判断坐标，写 needs_geocoding=true 并解释原因。",
                        "needs_geocoding": False,
                    }
                ],
                "routes": [
                    {
                        "id": "asset-route",
                        "name": "路线名称",
                        "summary": "路线叙事摘要",
                        "timeline": [
                            {
                                "year": 1,
                                "location": "地点",
                                "coordinates": [0, 0],
                                "event_type": "产地/加工/仓储/销售/传播",
                                "route": "该节点说明",
                            }
                        ],
                    }
                ],
                "timeline": [],
                "radar_metrics": [],
            },
            "brand_assets": {
                "slogan": "",
                "poster_copy": [],
                "whitepaper_outline": [],
                "dashboard_cards": [],
                "style_direction": [],
                "layout_direction": [],
            },
            "scores": {"completeness": 0, "evidence_strength": 0, "visualization_fit": 0, "risk": 0},
            "risks": [],
            "citations": [
                {
                    "id": "C1",
                    "claim": "被资产包使用的证据点",
                    "source": "文件名或 manual-note.txt",
                    "locator": "页码、表名、段落或 source 标记附近位置",
                    "quote": "不超过 40 字的原文摘录",
                    "confidence": 0.8,
                }
            ],
            "review_status": {
                "status": "pending",
                "summary": "人工复核状态摘要",
                "required_items": [
                    {"label": "产地证明", "status": "pending", "reason": "需要核验原件或编号"},
                    {"label": "检测指标", "status": "pending", "reason": "需要核验报告日期、机构和批次"},
                ],
            },
            "next_actions": [],
        },
    }
    return [
        {
            "role": "system",
            "content": (
                "你是 FlavorScape 的空间品牌资产分析器。请只输出合法 JSON。"
                "你的任务是把产品资料转为可执行的品牌资产包：海报、白皮书、大屏、地图节点、证据指标。"
                "visualization.map_nodes 和 routes.timeline 是后续地图真实渲染依据，不能泛泛描述；能判断地点时必须给 [经度, 纬度]。"
                "如果资料只给地点名但无法确定坐标，请仍保留节点名称，并写 needs_geocoding=true，不要编造精确坐标。"
                "每个关键证据尽量写入 citations，source 必须对应 [[source:文件名]] 标记，locator 写页码、表名、段落或附近标题。"
                "review_status 用来提示人工复核状态，不要声称已完成真实审核；默认 status 为 pending 或 needs_review。"
                "thinking_trace 只能写用户可见的处理步骤摘要，不要输出隐藏推理或长篇思维链。"
            ),
        },
        {
            "role": "user",
            "content": json.dumps(
                {
                    "required_schema": schema,
                    "context": {
                        "template_id": template_id,
                        "product_hint": product_hint,
                        "origin_hint": origin_hint,
                        "files": file_summary,
                    },
                    "document_text": text,
                },
                ensure_ascii=False,
            ),
        },
    ]


def _poster_chat_messages(payload: PosterChatRequest) -> list[dict[str, str]]:
    schema = {
        "reply": "给用户的简洁编辑建议",
        "suggestions": [
            {
                "field": "title|subtitle|poeticLine|narrative|theme|imagePosY|templateParams",
                "value": "可直接写入编辑器的值",
                "reason": "为什么这样改",
            }
        ],
        "style_options": [
            {"label": "样式名", "palette": "色系", "layout": "布局建议", "notes": "适用场景"}
        ],
    }
    return [
        {
            "role": "system",
            "content": (
                "你是 FlavorScape Studio 的海报编辑助手。请只输出合法 JSON。"
                "你可以帮助用户改标题、副标题、诗意短句、叙事正文、主题和图片焦点。"
                "theme 只支持 nature|heritage|indigo。templateParams 支持 layout、palette、density、imageShape、metricStyle、titleScale、badgeText。"
            ),
        },
        {
            "role": "user",
            "content": json.dumps(
                {
                    "required_schema": schema,
                    "current_poster": payload.poster,
                    "product": payload.product,
                    "conversation": [message.dict() for message in payload.messages],
                },
                ensure_ascii=False,
            ),
        },
    ]


def _studio_chat_messages(payload: StudioChatRequest) -> list[dict[str, str]]:
    schemas = {
        "poster": {
            "fields": ["title", "subtitle", "poeticLine", "narrative", "theme", "imagePosY", "templateParams"],
            "notes": "theme 只支持 nature|heritage|indigo；templateParams 可包含 layout、palette、density、imageShape、metricStyle、titleScale、badgeText。",
        },
        "archive": {
            "fields": ["title", "summary", "conclusion", "visibleEvidence"],
            "notes": "白皮书建议应聚焦摘要、结论、证据显示取舍和人工复核提示。",
        },
        "display": {
            "fields": ["title", "subtitle", "caption", "interactionMode", "selectedRouteId", "showTimeline"],
            "notes": "interactionMode 只支持 product|route；大屏建议应聚焦巡航叙事、节点展示和传播路径。",
        },
    }
    schema = {
        "reply": "给用户的简洁编辑建议",
        "suggestions": [
            {
                "field": "|".join(schemas[payload.output_type]["fields"]),
                "value": "可直接写入当前编辑器字段的值",
                "reason": "为什么这样改",
            }
        ],
        "style_options": [
            {"label": "样式名", "palette": "色系", "layout": "布局建议", "notes": "适用场景"}
        ],
        "next_actions": ["建议用户下一步执行的编辑动作"],
    }
    return [
        {
            "role": "system",
            "content": (
                "你是 FlavorScape Studio 的交付编辑助手。请只输出合法 JSON。"
                "你要根据当前产出类型提供可应用字段建议、样式方向和下一步动作。"
                f"当前产出类型是 {payload.output_type}。{schemas[payload.output_type]['notes']}"
                "不要假装已经完成真实检测或人工审核；涉及证据时提醒用户复核来源。"
            ),
        },
        {
            "role": "user",
            "content": json.dumps(
                {
                    "required_schema": schema,
                    "output_type": payload.output_type,
                    "current_output": payload.output,
                    "product": payload.product,
                    "asset_package": payload.asset_package,
                    "conversation": [message.dict() for message in payload.messages],
                },
                ensure_ascii=False,
            ),
        },
    ]


def _asset_chat_messages(payload: AssetChatRequest) -> list[dict[str, str]]:
    schema = {
        "reply": "给用户的业务建议，说明当前资产包下一步怎么推进",
        "recommendations": [
            {
                "target": "poster|archive|display|evidence|risk|workflow",
                "title": "建议标题",
                "detail": "具体执行建议",
                "priority": "high|medium|low",
            }
        ],
        "next_actions": ["下一步动作"],
    }
    return [
        {
            "role": "system",
            "content": (
                "你是 FlavorScape 的资产生产顾问。请只输出合法 JSON。"
                "你要帮助用户把分析结果推进成完整业务流程：资料补齐、证据审核、海报编辑、白皮书、大屏和交付。"
                "建议必须具体、可执行，并指出适合进入 Studio 的编辑动作。"
            ),
        },
        {
            "role": "user",
            "content": json.dumps(
                {
                    "required_schema": schema,
                    "asset_package": payload.asset_package,
                    "extraction": payload.extraction,
                    "conversation": [message.dict() for message in payload.messages],
                },
                ensure_ascii=False,
            ),
        },
    ]


def _call_deepseek_json(messages: list[dict[str, str]], max_tokens: int) -> dict[str, Any]:
    request_body = {
        "model": DEEPSEEK_MODEL,
        "messages": messages,
        "response_format": {"type": "json_object"},
        "temperature": 0.35,
        "max_tokens": max_tokens,
    }
    data = json.dumps(request_body, ensure_ascii=False).encode("utf-8")
    req = urllib.request.Request(
        f"{DEEPSEEK_BASE_URL.rstrip('/')}/chat/completions",
        data=data,
        headers={
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=120) as response:
        raw = response.read().decode("utf-8")
    envelope = json.loads(raw)
    content = envelope["choices"][0]["message"]["content"]
    try:
        return _parse_json_content(content)
    except ValueError:
        log.warning("DeepSeek returned non-JSON content; attempting JSON repair.")
        try:
            repaired = _repair_json_response(content, messages, max_tokens)
            return _parse_json_content(repaired)
        except Exception as exc:
            log.warning("DeepSeek JSON repair failed: %s", exc)
            return _fallback_from_text(content, messages)


def _repair_json_response(content: str, original_messages: list[dict[str, str]], max_tokens: int) -> str:
    repair_messages = [
        {
            "role": "system",
            "content": "你是 JSON 修复器。请把用户提供的文本转换为合法 JSON，只输出 JSON，不要 Markdown。",
        },
        {
            "role": "user",
            "content": json.dumps(
                {
                    "original_task": original_messages[0].get("content", "") if original_messages else "",
                    "non_json_response": content,
                },
                ensure_ascii=False,
            ),
        },
    ]
    request_body = {
        "model": DEEPSEEK_MODEL,
        "messages": repair_messages,
        "response_format": {"type": "json_object"},
        "temperature": 0,
        "max_tokens": min(max_tokens, 2200),
    }
    data = json.dumps(request_body, ensure_ascii=False).encode("utf-8")
    req = urllib.request.Request(
        f"{DEEPSEEK_BASE_URL.rstrip('/')}/chat/completions",
        data=data,
        headers={
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=120) as response:
        raw = response.read().decode("utf-8")
    envelope = json.loads(raw)
    return envelope["choices"][0]["message"]["content"]


def _require_deepseek() -> None:
    if not DEEPSEEK_API_KEY or DEEPSEEK_API_KEY == "your_deepseek_api_key_here":
        raise HTTPException(
            status_code=503,
            detail="DEEPSEEK_API_KEY is not configured. Add it to .env and restart the backend.",
        )


def _parse_json_content(content: str) -> dict[str, Any]:
    content = _strip_json_fence(content)
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        candidate = _extract_json_object(content)
        if candidate is None:
            raise ValueError("DeepSeek returned non-JSON content")
        return json.loads(candidate)


def _strip_json_fence(content: str) -> str:
    value = content.strip()
    fence = re.match(r"^```(?:json)?\s*(.*?)\s*```$", value, flags=re.S | re.I)
    return fence.group(1).strip() if fence else value


def _extract_json_object(content: str) -> str | None:
    start = content.find("{")
    end = content.rfind("}")
    if start == -1 or end == -1 or end <= start:
        return None
    return content[start:end + 1]


def _fallback_from_text(content: str, messages: list[dict[str, str]]) -> dict[str, Any]:
    text = re.sub(r"\s+", " ", content).strip()
    if _is_poster_chat(messages) or _is_studio_chat(messages):
        return {
            "reply": text[:1200] or "DeepSeek 返回了非 JSON 内容，暂未生成可应用字段。",
            "suggestions": [],
            "style_options": [],
            "next_actions": [],
        }

    summary = text[:700] or "DeepSeek 返回了非 JSON 内容，系统已生成兜底资产包。"
    return {
        "thinking_trace": [
            "DeepSeek 返回了非 JSON 文本，系统已保留其文本内容并生成兜底结构。",
            "建议补充更明确的产品名、产地、检测指标和目标渠道后重新分析。",
        ],
        "asset_package": {
            "product": {"name": "", "category": "", "origin": "", "ingredients": []},
            "evidence": {
                "lab_indicators": [],
                "certifications": [],
                "origin_claims": [],
                "process_steps": [],
            },
            "visualization": {
                "map_nodes": [],
                "routes": [],
                "timeline": [],
                "radar_metrics": [],
            },
            "brand_assets": {
                "slogan": summary,
                "poster_copy": [summary],
                "whitepaper_outline": ["请根据上传资料人工复核关键证据，并重新运行分析以获得结构化白皮书提纲。"],
                "dashboard_cards": ["可先把产地、工艺节点、检测指标整理为大屏卡片。"],
                "style_direction": ["根据文本语气选择自然绿、东方红或蓝印花主题。"],
                "layout_direction": ["当前为兜底结果，建议重新分析后再确定最终版式。"],
            },
            "scores": {
                "completeness": 35,
                "evidence_strength": 20,
                "visualization_fit": 30,
                "risk": 65,
            },
            "risks": ["模型未按 JSON 输出，结构化可信度较低。"],
            "citations": [],
            "review_status": {
                "status": "needs_review",
                "summary": "模型输出未能结构化，所有证据都需要人工复核。",
                "required_items": [
                    {"label": "模型原始回复", "status": "pending", "reason": "需人工判断是否可采信。"}
                ],
            },
            "next_actions": ["补充结构更清晰的资料，或在补充说明中列出产品、产地、证据和目标输出。"],
        },
    }


def _is_poster_chat(messages: list[dict[str, str]]) -> bool:
    if not messages:
        return False
    return "海报编辑助手" in messages[0].get("content", "")


def _is_studio_chat(messages: list[dict[str, str]]) -> bool:
    if not messages:
        return False
    return "交付编辑助手" in messages[0].get("content", "")


def _public_extraction(extraction: dict[str, Any]) -> dict[str, Any]:
    public = dict(extraction)
    public["files"] = [
        {key: value for key, value in item.items() if key != "text"}
        for item in extraction.get("files", [])
    ]
    return public


def _normalize_package(data: Any) -> dict[str, Any]:
    package = data if isinstance(data, dict) else {}
    return {
        "product": _with_defaults(package.get("product"), {
            "name": "",
            "category": "",
            "origin": "",
            "ingredients": [],
        }),
        "evidence": _with_defaults(package.get("evidence"), {
            "lab_indicators": [],
            "certifications": [],
            "origin_claims": [],
            "process_steps": [],
        }),
        "visualization": _with_defaults(package.get("visualization"), {
            "map_nodes": [],
            "routes": [],
            "timeline": [],
            "radar_metrics": [],
        }),
        "brand_assets": _with_defaults(package.get("brand_assets"), {
            "slogan": "",
            "poster_copy": [],
            "whitepaper_outline": [],
            "dashboard_cards": [],
            "style_direction": [],
            "layout_direction": [],
        }),
        "scores": _normalize_scores(package.get("scores")),
        "risks": _as_text_list(package.get("risks")),
        "citations": _normalize_citations(package.get("citations")),
        "review_status": _normalize_review_status(package.get("review_status") or package.get("manual_review")),
        "next_actions": _as_text_list(package.get("next_actions")),
    }


def _with_defaults(value: Any, defaults: dict[str, Any]) -> dict[str, Any]:
    merged = defaults.copy()
    if isinstance(value, dict):
        merged.update(value)
    return merged


def _normalize_scores(value: Any) -> dict[str, int]:
    defaults = {"completeness": 0, "evidence_strength": 0, "visualization_fit": 0, "risk": 0}
    if isinstance(value, dict):
        for key in defaults:
            defaults[key] = _clamp_score(value.get(key, 0))
    return defaults


def _normalize_suggestions(value: Any) -> list[dict[str, Any]]:
    allowed = {"title", "subtitle", "poeticLine", "narrative", "theme", "imagePosY", "templateParams"}
    suggestions = []
    for item in _as_list_of_dict(value):
        field = item.get("field")
        if field not in allowed:
            continue
        suggestions.append({
            "field": field,
            "value": item.get("value"),
            "reason": str(item.get("reason") or ""),
        })
    return suggestions[:8]


def _normalize_studio_suggestions(output_type: str, value: Any) -> list[dict[str, Any]]:
    allowed = {
        "poster": {"title", "subtitle", "poeticLine", "narrative", "theme", "imagePosY", "templateParams"},
        "archive": {"title", "summary", "conclusion", "visibleEvidence"},
        "display": {"title", "subtitle", "caption", "interactionMode", "selectedRouteId", "showTimeline"},
    }.get(output_type, set())
    suggestions = []
    for item in _as_list_of_dict(value):
        field = item.get("field")
        if field not in allowed:
            continue
        suggestions.append({
            "field": field,
            "value": item.get("value"),
            "reason": str(item.get("reason") or ""),
        })
    return suggestions[:10]


def _normalize_citations(value: Any) -> list[dict[str, Any]]:
    citations = []
    if isinstance(value, list):
        for index, item in enumerate(value, start=1):
            if isinstance(item, str):
                citations.append({
                    "id": f"C{index}",
                    "claim": item,
                    "source": "",
                    "locator": "",
                    "quote": "",
                    "confidence": 0,
                })
            elif isinstance(item, dict):
                citations.append({
                    "id": str(item.get("id") or f"C{index}"),
                    "claim": str(item.get("claim") or item.get("title") or item.get("label") or ""),
                    "source": str(item.get("source") or item.get("file") or ""),
                    "locator": str(item.get("locator") or item.get("location") or item.get("page") or ""),
                    "quote": str(item.get("quote") or item.get("excerpt") or "")[:180],
                    "confidence": _clamp_confidence(item.get("confidence")),
                })
    return citations[:16]


def _normalize_review_status(value: Any) -> dict[str, Any]:
    if not isinstance(value, dict):
        return {
            "status": "pending",
            "summary": "证据尚未人工复核。",
            "required_items": [],
        }
    status = str(value.get("status") or "pending")
    if status not in {"pending", "needs_review", "in_review", "approved", "rejected"}:
        status = "pending"
    return {
        "status": status,
        "summary": str(value.get("summary") or "证据尚未人工复核。"),
        "required_items": _normalize_review_items(value.get("required_items") or value.get("items")),
    }


def _normalize_review_items(value: Any) -> list[dict[str, str]]:
    items = []
    if isinstance(value, list):
        for item in value:
            if isinstance(item, str):
                items.append({"label": item, "status": "pending", "reason": ""})
            elif isinstance(item, dict):
                status = str(item.get("status") or "pending")
                if status not in {"pending", "needs_review", "in_review", "approved", "rejected"}:
                    status = "pending"
                items.append({
                    "label": str(item.get("label") or item.get("title") or item.get("name") or "复核项"),
                    "status": status,
                    "reason": str(item.get("reason") or item.get("detail") or ""),
                })
    return items[:12]


def _as_list_of_dict(value: Any) -> list[dict[str, Any]]:
    if not isinstance(value, list):
        return []
    return [item for item in value if isinstance(item, dict)]


def _as_text_list(value: Any) -> list[str]:
    if not isinstance(value, list):
        return []
    result = []
    for item in value:
        if isinstance(item, str):
            result.append(item)
        elif isinstance(item, dict):
            result.append(item.get("title") or item.get("name") or item.get("label") or item.get("description") or json.dumps(item, ensure_ascii=False))
        else:
            result.append(str(item))
    return [item for item in result if item]


def _clamp_score(value: Any) -> int:
    try:
        return max(0, min(100, round(float(value))))
    except (TypeError, ValueError):
        return 0


def _clamp_confidence(value: Any) -> float:
    try:
        numeric = float(value)
        if numeric > 1:
            numeric = numeric / 100
        return round(max(0, min(1, numeric)), 2)
    except (TypeError, ValueError):
        return 0
